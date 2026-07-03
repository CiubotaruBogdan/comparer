#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document Comparator v1.3
Compară documente la nivel de paragraf (PDF cu PDF, DOCX cu DOCX),
cu highlighting al diferențelor, statistici detaliate,
export PDF/HTML, și vizualizare optimizată pentru documente mari.
"""

import sys
import difflib
import html as html_lib
from pathlib import Path
from typing import Optional
from datetime import datetime

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph as DocxParagraph
import fitz  # PyMuPDF

if sys.platform == "win32":
    import ctypes
    ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("Comparer.1.3")

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QSplitter, QFrame,
    QFileDialog, QStackedWidget, QMessageBox, QSizePolicy, QProgressBar,
    QShortcut, QScrollArea, QAbstractScrollArea, QStyleOptionViewItem,
    QStyledItemDelegate, QListView,
)
from PyQt5.QtCore import (
    Qt, pyqtSignal, QThread, pyqtSlot, QTimer, QAbstractListModel,
    QModelIndex, QSize, QRect, QPoint, QRectF,
)
from PyQt5.QtGui import (
    QFont, QColor, QPalette, QIcon, QKeySequence, QPainter,
    QTextDocument, QAbstractTextDocumentLayout, QPen, QBrush,
    QFontMetrics,
)
from PyQt5.QtPrintSupport import QPrinter


def _resource(relative: str) -> Path:
    if hasattr(sys, "_MEIPASS"):
        return Path(sys._MEIPASS) / relative
    return Path(__file__).parent / relative


def _load_icon() -> QIcon:
    for name in ("icon.png", "icon.ico"):
        p = _resource(name)
        if p.exists():
            return QIcon(str(p))
    return QIcon()


# ══════════════════════════════════════════════════════════════════════════════
# 1. Document Readers
# ══════════════════════════════════════════════════════════════════════════════

def _iter_block_texts(element, doc: DocxDocument) -> list:
    texts = []
    for child in element.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            text = DocxParagraph(child, doc).text.strip()
            if text:
                texts.append(text)
        elif tag == "tbl":
            for tr in child.findall(qn("w:tr")):
                for tc in tr.findall(qn("w:tc")):
                    texts.extend(_iter_block_texts(tc, doc))
    return texts


def _iter_header_footer_texts(doc: DocxDocument) -> list:
    """Extract text from all headers and footers in the document (deduplicated)."""
    seen = set()
    texts = []
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header is not None and not header.is_linked_to_previous:
                for para in header.paragraphs:
                    text = para.text.strip()
                    if text and text not in seen:
                        seen.add(text)
                        texts.append(f"[HEADER] {text}")
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer is not None and not footer.is_linked_to_previous:
                for para in footer.paragraphs:
                    text = para.text.strip()
                    if text and text not in seen:
                        seen.add(text)
                        texts.append(f"[FOOTER] {text}")
    # Fallback: if all are linked (single section doc), read the main header/footer
    if not texts:
        for section in doc.sections:
            for header in (section.header,):
                if header is not None:
                    for para in header.paragraphs:
                        text = para.text.strip()
                        if text and text not in seen:
                            seen.add(text)
                            texts.append(f"[HEADER] {text}")
            for footer in (section.footer,):
                if footer is not None:
                    for para in footer.paragraphs:
                        text = para.text.strip()
                        if text and text not in seen:
                            seen.add(text)
                            texts.append(f"[FOOTER] {text}")
    return texts


def read_docx(path: str) -> list:
    doc = DocxDocument(path)
    header_footer_texts = _iter_header_footer_texts(doc)
    body_texts = _iter_block_texts(doc.element.body, doc)
    return header_footer_texts + body_texts


def read_pdf(path: str) -> list:
    paragraphs = []
    with fitz.open(path) as doc:
        for page in doc:
            blocks = page.get_text("blocks")
            blocks_sorted = sorted(blocks, key=lambda b: (round(b[1] / 12) * 12, b[0]))
            for block in blocks_sorted:
                if len(block) > 6 and block[6] == 0:
                    text = " ".join(block[4].split())
                    if len(text) > 3:
                        paragraphs.append(text)
    return paragraphs


def read_document(path: str) -> list:
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        return read_docx(path)
    elif ext == ".pdf":
        return read_pdf(path)
    else:
        raise ValueError(f"Format nesuportat: {ext}\nSunt acceptate: .docx, .pdf")


def validate_same_type(path1: str, path2: str) -> Optional[str]:
    """Validate that both files have the same extension. Returns error message or None."""
    ext1 = Path(path1).suffix.lower()
    ext2 = Path(path2).suffix.lower()
    if ext1 != ext2:
        return (
            f"Tipuri de fisiere diferite!\n\n"
            f"Document 1: {ext1.upper().lstrip('.')}\n"
            f"Document 2: {ext2.upper().lstrip('.')}\n\n"
            f"Comparatia este permisa doar intre fisiere de acelasi tip:\n"
            f"  \u2022 PDF cu PDF\n"
            f"  \u2022 DOCX cu DOCX"
        )
    return None


# ══════════════════════════════════════════════════════════════════════════════
# 2. Motor de comparație
# ══════════════════════════════════════════════════════════════════════════════

def word_level_diff(text1: str, text2: str):
    words1 = text1.split()
    words2 = text2.split()
    matcher = difflib.SequenceMatcher(None, words1, words2, autojunk=False)

    left_parts = []
    right_parts = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        w1 = html_lib.escape(" ".join(words1[i1:i2]))
        w2 = html_lib.escape(" ".join(words2[j1:j2]))

        if tag == "equal":
            left_parts.append(w1)
            right_parts.append(w1)
        elif tag == "replace":
            left_parts.append(
                f'<span style="background:#FFCDD2;text-decoration:line-through;'
                f'border-radius:2px;padding:0 2px;">{w1}</span>'
            )
            right_parts.append(
                f'<span style="background:#C8E6C9;border-radius:2px;padding:0 2px;">{w2}</span>'
            )
        elif tag == "delete":
            left_parts.append(
                f'<span style="background:#FFCDD2;text-decoration:line-through;'
                f'border-radius:2px;padding:0 2px;">{w1}</span>'
            )
        elif tag == "insert":
            right_parts.append(
                f'<span style="background:#C8E6C9;border-radius:2px;padding:0 2px;">{w2}</span>'
            )

    return " ".join(left_parts), " ".join(right_parts)


def compare_documents(paras1: list, paras2: list) -> dict:
    matcher = difflib.SequenceMatcher(None, paras1, paras2, autojunk=False)
    similarity = matcher.ratio()

    left_blocks = []
    right_blocks = []

    stats = {
        "total_left":  len(paras1),
        "total_right": len(paras2),
        "equal":    0,
        "modified": 0,
        "added":    0,
        "deleted":  0,
    }

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for k in range(i2 - i1):
                left_blocks.append(("equal", paras1[i1 + k], False))
                right_blocks.append(("equal", paras2[j1 + k], False))
                stats["equal"] += 1

        elif tag == "replace":
            lp = paras1[i1:i2]
            rp = paras2[j1:j2]
            common = min(len(lp), len(rp))
            stats["modified"] += common
            stats["deleted"]  += len(lp) - common
            stats["added"]    += len(rp) - common

            for k in range(max(len(lp), len(rp))):
                has_l = k < len(lp)
                has_r = k < len(rp)
                if has_l and has_r:
                    lh, rh = word_level_diff(lp[k], rp[k])
                    left_blocks.append(("replace", lh, True))
                    right_blocks.append(("replace", rh, True))
                elif has_l:
                    left_blocks.append(("delete", lp[k], False))
                    right_blocks.append(("placeholder", "", False))
                else:
                    left_blocks.append(("placeholder", "", False))
                    right_blocks.append(("insert", rp[k], False))

        elif tag == "delete":
            for k in range(i2 - i1):
                left_blocks.append(("delete", paras1[i1 + k], False))
                right_blocks.append(("placeholder", "", False))
                stats["deleted"] += 1

        elif tag == "insert":
            for k in range(j2 - j1):
                left_blocks.append(("placeholder", "", False))
                right_blocks.append(("insert", paras2[j1 + k], False))
                stats["added"] += 1

    diff_positions = []
    in_diff = False
    for i in range(len(left_blocks)):
        ls = left_blocks[i][0]
        rs = right_blocks[i][0]
        is_diff = ls not in ("equal", "placeholder") or rs not in ("equal", "placeholder")
        if is_diff and not in_diff:
            diff_positions.append(i)
        in_diff = is_diff

    return {
        "similarity":    similarity,
        "left_blocks":   left_blocks,
        "right_blocks":  right_blocks,
        "diff_positions": diff_positions,
        "stats":         stats,
        "diff_count":    len(diff_positions),
    }


def render_html(
    blocks: list,
    side: str,
    diff_positions: list = None,
    padding_y: int = 7,
    interactive: bool = False,
) -> str:
    """Generates HTML for export reports only."""
    STYLE = {
        "equal":       ("#FFFFFF", "#E0E0E0", "3px solid"),
        "replace":     ("#FFECB3", "#FFB300", "3px solid"),
        "delete":      ("#FFCDD2", "#E53935", "3px solid"),
        "insert":      ("#C8E6C9", "#43A047", "3px solid"),
        "placeholder": ("#F9F9F9", "#EEEEEE", "3px dotted"),
    }
    BADGE = {
        "delete": ('<span style="font-size:9px;background:#E53935;color:#fff;'
                   'padding:1px 6px;border-radius:2px;margin-right:6px;">STERS</span>'),
        "insert": ('<span style="font-size:9px;background:#43A047;color:#fff;'
                   'padding:1px 6px;border-radius:2px;margin-right:6px;">NOU</span>'),
    }

    anchor_map: dict = {}
    if diff_positions:
        for group_idx, block_idx in enumerate(diff_positions):
            anchor_map[block_idx] = group_idx

    lines = [
        '<!DOCTYPE html><html>',
        '<head><meta charset="utf-8"></head>',
        '<body style="font-family:system-ui,Arial,sans-serif;font-size:13px;'
        'line-height:1.55;margin:0;padding:6px;background:#FAFAFA;color:#111;">',
    ]

    for block_idx, (status, content, is_html) in enumerate(blocks):
        bg, bl_color, bl_width = STYLE.get(status, STYLE["equal"])

        anchor = ""
        if block_idx in anchor_map:
            anchor = f'<a name="diff-{anchor_map[block_idx]}"></a>'

        badge = ""
        if status == "delete" and side == "left":
            badge = BADGE["delete"]
        elif status == "insert" and side == "right":
            badge = BADGE["insert"]

        if status == "placeholder":
            inner = ""
        elif is_html:
            inner = f"{badge}<span>{content}</span>"
        else:
            inner = f"{badge}<span>{html_lib.escape(content)}</span>"

        body = f"{anchor}{inner}"
        lines.append(
            f'<div style="background:{bg};border-left:{bl_width} {bl_color};'
            f'margin:2px 0;padding:{padding_y}px 10px;'
            f'border-radius:0 2px 2px 0;min-height:22px;overflow:hidden;">'
            f'{body}</div>'
        )

    lines.append('</body></html>')
    return "\n".join(lines)


def export_html_report(path1: str, path2: str, result: dict) -> str:
    sim   = result["similarity"]
    stats = result["stats"]
    n1 = html_lib.escape(Path(path1).name)
    n2 = html_lib.escape(Path(path2).name)
    export_l = render_html(result["left_blocks"],  "left",  result["diff_positions"])
    export_r = render_html(result["right_blocks"], "right", result["diff_positions"])
    body_l = export_l.split("<body", 1)[1].split(">", 1)[1].rsplit("</body>", 1)[0]
    body_r = export_r.split("<body", 1)[1].split(">", 1)[1].rsplit("</body>", 1)[0]

    return f"""<!DOCTYPE html>
<html lang="ro">
<head>
<meta charset="utf-8">
<title>Comparatie — {n1} vs {n2}</title>
<style>
  *   {{ box-sizing:border-box; margin:0; padding:0; }}
  body {{ font-family:system-ui,Arial,sans-serif; background:#fff; color:#111; }}
  .hdr {{ background:#111; color:#fff; padding:14px 24px; }}
  .hdr h1 {{ font-size:16px; font-weight:600; }}
  .hdr p  {{ font-size:11px; color:#aaa; margin-top:3px; }}
  .stats  {{ background:#F4F4F4; border-bottom:1px solid #ddd;
             padding:10px 24px; display:flex; gap:28px; align-items:center; flex-wrap:wrap; }}
  .stat .val {{ font-size:20px; font-weight:700; }}
  .stat .lbl {{ font-size:9px; color:#777; letter-spacing:.5px; }}
  .sim        {{ font-size:32px; font-weight:700; }}
  .cols       {{ display:grid; grid-template-columns:1fr 1fr; }}
  .col-hdr    {{ background:#222; color:#fff; padding:7px 12px;
                font-size:11px; font-weight:600; letter-spacing:.3px; }}
  .col-body   {{ padding:6px; background:#fff; min-height:80vh;
                border-right:1px solid #ddd; }}
</style>
</head>
<body>
<div class="hdr">
  <h1>Raport Comparatie Documente</h1>
  <p>{n1} &nbsp;&#x27F7;&nbsp; {n2}</p>
</div>
<div class="stats">
  <div class="stat"><div class="sim">{sim*100:.1f}%</div><div class="lbl">SIMILARITATE</div></div>
  <div style="width:1px;background:#ccc;height:36px;"></div>
  <div class="stat"><div class="val">{stats["equal"]}</div><div class="lbl">IDENTICE</div></div>
  <div class="stat"><div class="val">{stats["modified"]}</div><div class="lbl">MODIFICATE</div></div>
  <div class="stat"><div class="val">{stats["added"]}</div><div class="lbl">ADAUGATE</div></div>
  <div class="stat"><div class="val">{stats["deleted"]}</div><div class="lbl">STERSE</div></div>
  <div class="stat"><div class="val">{stats["total_left"]}</div><div class="lbl">PAR. DOC 1</div></div>
  <div class="stat"><div class="val">{stats["total_right"]}</div><div class="lbl">PAR. DOC 2</div></div>
</div>
<div class="cols">
  <div>
    <div class="col-hdr">Document 1 — Baza: {n1}</div>
    <div class="col-body">{body_l}</div>
  </div>
  <div>
    <div class="col-hdr">Document 2 — Comparatie: {n2}</div>
    <div class="col-body">{body_r}</div>
  </div>
</div>
</body></html>"""


def export_pdf_report(path1: str, path2: str, result: dict) -> str:
    """Generate HTML content optimized for PDF rendering via QPrinter."""
    sim = result["similarity"]
    stats = result["stats"]
    n1 = html_lib.escape(Path(path1).name)
    n2 = html_lib.escape(Path(path2).name)
    now = datetime.now().strftime("%d.%m.%Y  %H:%M:%S")

    # Build left/right block rows for the table
    left_blocks = result["left_blocks"]
    right_blocks = result["right_blocks"]

    STYLE = {
        "equal":       ("#FFFFFF", "#E0E0E0"),
        "replace":     ("#FFECB3", "#FFB300"),
        "delete":      ("#FFCDD2", "#E53935"),
        "insert":      ("#C8E6C9", "#43A047"),
        "placeholder": ("#F9F9F9", "#EEEEEE"),
    }
    BADGE = {
        "delete": '<span style="font-size:8px;background:#E53935;color:#fff;padding:1px 4px;border-radius:2px;margin-right:4px;">STERS</span>',
        "insert": '<span style="font-size:8px;background:#43A047;color:#fff;padding:1px 4px;border-radius:2px;margin-right:4px;">NOU</span>',
    }

    rows = []
    for idx in range(len(left_blocks)):
        l_status, l_content, l_is_html = left_blocks[idx]
        r_status, r_content, r_is_html = right_blocks[idx]

        l_bg, l_bc = STYLE.get(l_status, STYLE["equal"])
        r_bg, r_bc = STYLE.get(r_status, STYLE["equal"])

        # Left cell content
        l_badge = ""
        if l_status == "delete":
            l_badge = BADGE["delete"]
        l_text = ""
        if l_status != "placeholder":
            l_text = l_content if l_is_html else html_lib.escape(l_content)

        # Right cell content
        r_badge = ""
        if r_status == "insert":
            r_badge = BADGE["insert"]
        r_text = ""
        if r_status != "placeholder":
            r_text = r_content if r_is_html else html_lib.escape(r_content)

        rows.append(
            f'<tr>'
            f'<td style="background:{l_bg};border-left:3px solid {l_bc};padding:4px 6px;'
            f'vertical-align:top;width:50%;font-size:9px;line-height:1.4;">'
            f'{l_badge}{l_text}</td>'
            f'<td style="background:{r_bg};border-left:3px solid {r_bc};padding:4px 6px;'
            f'vertical-align:top;width:50%;font-size:9px;line-height:1.4;">'
            f'{r_badge}{r_text}</td>'
            f'</tr>'
        )

    table_rows = "\n".join(rows)

    return f"""<!DOCTYPE html>
<html lang="ro">
<head>
<meta charset="utf-8">
<style>
  * {{ box-sizing:border-box; margin:0; padding:0; }}
  body {{ font-family: Arial, sans-serif; font-size:10px; color:#111; }}
  .title {{ font-size:16px; font-weight:bold; text-align:center; margin-top:12px; }}
  .subtitle {{ font-size:10px; color:#555; text-align:center; margin-top:6px; margin-bottom:14px; }}
  .stats-table {{ width:100%; border-collapse:collapse; margin-bottom:10px; }}
  .stats-table td {{ padding:6px 10px; text-align:center; border:1px solid #ddd; }}
  .stats-table .val {{ font-size:14px; font-weight:bold; }}
  .stats-table .lbl {{ font-size:8px; color:#777; letter-spacing:0.5px; }}
  .compare-table {{ width:100%; border-collapse:collapse; }}
  .compare-table td {{ border-bottom:1px solid #eee; }}
  .col-hdr {{ background:#222; color:#fff; padding:5px 8px; font-size:9px; font-weight:bold; }}
</style>
</head>
<body>
<div class="title">Documente comparate: {n1}, {n2}</div>
<div class="subtitle">Data raport: {now}</div>

<table class="stats-table">
<tr>
  <td><div class="val">{sim*100:.1f}%</div><div class="lbl">SIMILARITATE</div></td>
  <td><div class="val">{stats["equal"]}</div><div class="lbl">IDENTICE</div></td>
  <td><div class="val">{stats["modified"]}</div><div class="lbl">MODIFICATE</div></td>
  <td><div class="val">{stats["added"]}</div><div class="lbl">ADAUGATE</div></td>
  <td><div class="val">{stats["deleted"]}</div><div class="lbl">STERSE</div></td>
  <td><div class="val">{stats["total_left"]}</div><div class="lbl">PAR. DOC 1</div></td>
  <td><div class="val">{stats["total_right"]}</div><div class="lbl">PAR. DOC 2</div></td>
</tr>
</table>

<table class="compare-table">
<tr>
  <td class="col-hdr" style="width:50%;">Document 1 — Baza: {n1}</td>
  <td class="col-hdr" style="width:50%;">Document 2 — Comparatie: {n2}</td>
</tr>
{table_rows}
</table>
</body></html>"""


# ══════════════════════════════════════════════════════════════════════════════
# 3. Thread de lucru
# ══════════════════════════════════════════════════════════════════════════════

class CompareWorker(QThread):
    finished = pyqtSignal(dict)
    error    = pyqtSignal(str)
    progress = pyqtSignal(str)

    def __init__(self, path1: str, path2: str):
        super().__init__()
        self.path1 = path1
        self.path2 = path2

    def run(self):
        try:
            self.progress.emit("Se citeste documentul 1...")
            paras1 = read_document(self.path1)
            self.progress.emit("Se citeste documentul 2...")
            paras2 = read_document(self.path2)

            if not paras1:
                self.error.emit(
                    "Nu s-au putut extrage paragrafe din Document 1.\n"
                    "Verificati ca fisierul nu este scanat (imagine)."
                )
                return
            if not paras2:
                self.error.emit(
                    "Nu s-au putut extrage paragrafe din Document 2.\n"
                    "Verificati ca fisierul nu este scanat (imagine)."
                )
                return

            self.progress.emit("Se compara documentele...")
            result = compare_documents(paras1, paras2)
            result["path1"] = self.path1
            result["path2"] = self.path2
            self.finished.emit(result)
        except Exception as e:
            self.error.emit(str(e))


# ══════════════════════════════════════════════════════════════════════════════
# 4. DropZone widget
# ══════════════════════════════════════════════════════════════════════════════

class DropZone(QFrame):
    file_loaded = pyqtSignal(str)

    def __init__(self, title: str, parent=None):
        super().__init__(parent)
        self.file_path: Optional[str] = None
        self._title = title
        self._last_dir = ""
        self.setAcceptDrops(True)
        self.setMinimumSize(280, 200)
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self._setup_ui()
        self._apply_style(False)

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setAlignment(Qt.AlignCenter)
        layout.setSpacing(6)
        layout.setContentsMargins(20, 16, 20, 16)

        self.title_lbl = QLabel(self._title)
        self.title_lbl.setAlignment(Qt.AlignCenter)
        f = QFont()
        f.setPointSize(10)
        f.setWeight(QFont.DemiBold)
        self.title_lbl.setFont(f)
        self.title_lbl.setStyleSheet("color:#111; background:transparent;")

        self.hint_lbl = QLabel("drag & drop  sau  click pentru a selecta")
        self.hint_lbl.setAlignment(Qt.AlignCenter)
        f2 = QFont()
        f2.setPointSize(8)
        self.hint_lbl.setFont(f2)
        self.hint_lbl.setStyleSheet("color:#BBB; background:transparent;")

        self.file_lbl = QLabel("")
        self.file_lbl.setAlignment(Qt.AlignCenter)
        f3 = QFont()
        f3.setPointSize(9)
        self.file_lbl.setFont(f3)
        self.file_lbl.setStyleSheet("color:#333; background:transparent;")
        self.file_lbl.setWordWrap(True)

        # Clear button (X) — hidden until a file is loaded
        self.clear_btn = QPushButton("\u2715")
        self.clear_btn.setFixedSize(22, 22)
        self.clear_btn.setCursor(Qt.PointingHandCursor)
        self.clear_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                border: none;
                color: #999;
                font-size: 14px;
                font-weight: bold;
                border-radius: 11px;
            }
            QPushButton:hover {
                background: #E0E0E0;
                color: #333;
            }
            QPushButton:pressed {
                background: #CCC;
            }
        """)
        self.clear_btn.hide()
        self.clear_btn.clicked.connect(self._clear_file)

        # Top row with clear button in top-right corner
        top_row = QHBoxLayout()
        top_row.setContentsMargins(0, 0, 0, 0)
        top_row.addStretch()
        top_row.addWidget(self.clear_btn)

        layout.addLayout(top_row)
        layout.addWidget(self.title_lbl)
        layout.addSpacing(2)
        layout.addWidget(self.hint_lbl)
        layout.addWidget(self.file_lbl)

    def _apply_style(self, loaded: bool):
        if loaded:
            self.setStyleSheet("""
                DropZone {
                    border: 1px solid #111;
                    border-radius: 6px;
                    background: #fff;
                }
                DropZone QLabel {
                    background: transparent;
                }
            """)
        else:
            self.setStyleSheet("""
                DropZone {
                    border: 1px dashed #BBBBBB;
                    border-radius: 6px;
                    background: #FAFAFA;
                }
                DropZone:hover {
                    border-color: #555;
                    background: #F4F4F4;
                }
                DropZone QLabel {
                    background: transparent;
                }
            """)

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self._browse()

    def _browse(self):
        path, _ = QFileDialog.getOpenFileName(
            self, f"Selecteaza {self._title}",
            self._last_dir, "Documente (*.docx *.pdf);;Word (*.docx);;PDF (*.pdf)"
        )
        if path:
            self._set_file(path)

    def _clear_file(self):
        """Clear the loaded file and reset the DropZone."""
        self.reset()
        self.file_loaded.emit("")  # Signal with empty path to update button state

    def _set_file_display(self, path: str):
        self.file_path = path
        self._last_dir = str(Path(path).parent)
        self.hint_lbl.setText("incarcat")
        self.hint_lbl.setStyleSheet("color:#AAA; background:transparent;")
        self.file_lbl.setText(Path(path).name)
        self.clear_btn.show()
        self._apply_style(True)

    def _set_file(self, path: str):
        self._set_file_display(path)
        self.file_loaded.emit(path)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            url = event.mimeData().urls()[0]
            if url.toLocalFile().lower().endswith((".docx", ".pdf")):
                event.acceptProposedAction()
                self.setStyleSheet("""
                    DropZone {
                        border: 1px solid #111;
                        border-radius: 6px;
                        background: #EFEFEF;
                    }
                    DropZone QLabel {
                        background: transparent;
                    }
                """)

    def dragLeaveEvent(self, event):
        self._apply_style(self.file_path is not None)

    def dropEvent(self, event):
        urls = event.mimeData().urls()
        if urls:
            path = urls[0].toLocalFile()
            if path.lower().endswith((".docx", ".pdf")):
                self._set_file(path)

    def reset(self):
        self.file_path = None
        self.hint_lbl.setText("drag & drop  sau  click pentru a selecta")
        self.hint_lbl.setStyleSheet("color:#BBB; background:transparent;")
        self.file_lbl.setText("")
        self.clear_btn.hide()
        self._apply_style(False)


# ══════════════════════════════════════════════════════════════════════════════
# 5. Virtualized ComparePanel (memory-efficient for large documents)
# ══════════════════════════════════════════════════════════════════════════════

class BlockListModel(QAbstractListModel):
    """
    Model that holds block data (status, content, is_html, side) without
    creating any widgets. Only raw data is stored — rendering is done by
    the delegate on demand for visible items only.
    """

    def __init__(self, parent=None):
        super().__init__(parent)
        self._blocks: list = []   # list of (status, content, is_html)
        self._side: str = "left"
        self._selected_idx: int = -1

    def set_data(self, blocks: list, side: str):
        self.beginResetModel()
        self._blocks = blocks
        self._side = side
        self._selected_idx = -1
        self.endResetModel()

    def rowCount(self, parent=QModelIndex()):
        return len(self._blocks)

    def data(self, index, role=Qt.DisplayRole):
        if not index.isValid():
            return None
        row = index.row()
        if row < 0 or row >= len(self._blocks):
            return None

        if role == Qt.UserRole:
            # Return full block tuple + side + selected
            status, content, is_html = self._blocks[row]
            return (status, content, is_html, self._side, row == self._selected_idx)
        return None

    def set_selected(self, idx: int):
        old = self._selected_idx
        self._selected_idx = idx
        if old >= 0 and old < len(self._blocks):
            i = self.index(old)
            self.dataChanged.emit(i, i, [Qt.UserRole])
        if idx >= 0 and idx < len(self._blocks):
            i = self.index(idx)
            self.dataChanged.emit(i, i, [Qt.UserRole])

    def block_count(self) -> int:
        return len(self._blocks)


class BlockDelegate(QStyledItemDelegate):
    """
    Custom delegate that paints paragraph blocks using QPainter directly.
    No widgets are created — this is the key to memory efficiency.
    Only visible items are rendered by the QListView viewport.
    """

    _BG = {
        "equal":       QColor("#FFFFFF"),
        "replace":     QColor("#FFECB3"),
        "delete":      QColor("#FFCDD2"),
        "insert":      QColor("#C8E6C9"),
        "placeholder": QColor("#F9F9F9"),
    }
    _BORDER_COLOR = {
        "equal":       QColor("#E0E0E0"),
        "replace":     QColor("#FFB300"),
        "delete":      QColor("#E53935"),
        "insert":      QColor("#43A047"),
        "placeholder": QColor("#EEEEEE"),
    }
    _SEL_BG = QColor("#EBEBEB")
    _SEL_BORDER = QColor("#888888")

    def __init__(self, parent=None):
        super().__init__(parent)
        self._font = QFont()
        self._font.setPointSize(10)
        self._doc_cache = {}  # LRU-like cache for QTextDocument (limited size)
        self._cache_max = 200  # max cached documents

    def _get_doc(self, html_content: str, width: int) -> QTextDocument:
        """Get or create a QTextDocument for the given content and width."""
        key = (html_content, width)
        if key in self._doc_cache:
            return self._doc_cache[key]

        doc = QTextDocument()
        doc.setDefaultFont(self._font)
        doc.setTextWidth(width)
        doc.setHtml(html_content)

        # Evict oldest if cache full
        if len(self._doc_cache) >= self._cache_max:
            # Remove first item (oldest)
            first_key = next(iter(self._doc_cache))
            del self._doc_cache[first_key]

        self._doc_cache[key] = doc
        return doc

    def _build_html(self, status: str, content: str, is_html: bool, side: str) -> str:
        """Build the rich text HTML for a block."""
        if status == "placeholder":
            return ""

        badge = ""
        if status == "delete" and side == "left":
            badge = ('<span style="font-size:8px;background:#E53935;color:#fff;'
                     'padding:1px 5px;border-radius:2px;margin-right:6px;">STERS</span> ')
        elif status == "insert" and side == "right":
            badge = ('<span style="font-size:8px;background:#43A047;color:#fff;'
                     'padding:1px 5px;border-radius:2px;margin-right:6px;">NOU</span> ')

        text = content if is_html else html_lib.escape(content)
        return badge + text

    def paint(self, painter: QPainter, option: QStyleOptionViewItem, index: QModelIndex):
        data = index.data(Qt.UserRole)
        if data is None:
            return

        status, content, is_html, side, selected = data
        rect = option.rect

        painter.save()
        painter.setClipRect(rect)

        # Background
        bg = self._SEL_BG if selected else self._BG.get(status, self._BG["equal"])
        painter.fillRect(rect, bg)

        # Left border (3px)
        border_color = self._SEL_BORDER if selected else self._BORDER_COLOR.get(status, self._BORDER_COLOR["equal"])
        border_rect = QRect(rect.left(), rect.top(), 3, rect.height())
        painter.fillRect(border_rect, border_color)

        # Content
        if status != "placeholder" and content:
            html_content = self._build_html(status, content, is_html, side)
            content_width = rect.width() - 20  # 10px padding on each side
            doc = self._get_doc(html_content, content_width)

            painter.translate(rect.left() + 13, rect.top() + 7)
            doc.drawContents(painter)

        painter.restore()

    def sizeHint(self, option: QStyleOptionViewItem, index: QModelIndex) -> QSize:
        data = index.data(Qt.UserRole)
        if data is None:
            return QSize(100, 32)

        status, content, is_html, side, selected = data

        if status == "placeholder" or not content:
            return QSize(option.rect.width() if option.rect.width() > 0 else 100, 32)

        html_content = self._build_html(status, content, is_html, side)
        width = option.rect.width() - 20 if option.rect.width() > 20 else 400
        doc = self._get_doc(html_content, width)
        h = int(doc.size().height()) + 14  # 7px top + 7px bottom padding
        return QSize(option.rect.width() if option.rect.width() > 0 else 100, max(h, 32))

    def clear_cache(self):
        """Clear the document cache (call when loading new data)."""
        self._doc_cache.clear()


class ComparePanel(QListView):
    """
    Memory-efficient panel using QListView with a custom model and delegate.
    Only visible items are rendered — no widget is created per paragraph.
    This drastically reduces memory usage for large documents (thousands of paragraphs).
    """
    block_clicked = pyqtSignal(int)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.setStyleSheet("""
            QListView {
                border: none;
                background: #FAFAFA;
            }
            QListView::item {
                margin: 1px 0px;
            }
        """)
        self.setSelectionMode(QListView.NoSelection)
        self.setEditTriggers(QListView.NoEditTriggers)
        self.setUniformItemSizes(False)
        self.setWordWrap(True)
        self.setSpacing(1)

        self._model = BlockListModel(self)
        self._delegate = BlockDelegate(self)
        self.setModel(self._model)
        self.setItemDelegate(self._delegate)

        self.clicked.connect(self._on_item_clicked)

        self._partner: Optional['ComparePanel'] = None
        self._row_heights: list = []

    def load_blocks(self, blocks: list, side: str, diff_positions: list):
        self._delegate.clear_cache()
        self._model.set_data(blocks, side)
        self._row_heights = []
        # Schedule height computation after layout
        QTimer.singleShot(50, self._compute_row_heights)

    def _compute_row_heights(self):
        """Pre-compute row heights for scroll synchronization."""
        n = self._model.block_count()
        vp_width = self.viewport().width()
        self._row_heights = []
        option = QStyleOptionViewItem()
        option.rect = QRect(0, 0, vp_width, 0)
        for i in range(n):
            idx = self._model.index(i)
            size = self._delegate.sizeHint(option, idx)
            self._row_heights.append(size.height())

    def set_selected(self, idx: int):
        self._model.set_selected(idx)

    def scroll_to_block(self, idx: int):
        if 0 <= idx < self._model.block_count():
            index = self._model.index(idx)
            self.scrollTo(index, QListView.PositionAtTop)

    def get_block_count(self) -> int:
        return self._model.block_count()

    def _on_item_clicked(self, index: QModelIndex):
        if index.isValid():
            data = index.data(Qt.UserRole)
            if data and data[0] != "placeholder":
                self.block_clicked.emit(index.row())


# ══════════════════════════════════════════════════════════════════════════════
# 6. Fereastra principala
# ══════════════════════════════════════════════════════════════════════════════

def _flat_btn(text: str, dark: bool = False) -> QPushButton:
    btn = QPushButton(text)
    f = QFont()
    f.setPointSize(9)
    btn.setFont(f)
    btn.setCursor(Qt.PointingHandCursor)
    if dark:
        btn.setStyleSheet("""
            QPushButton {
                background:#111; color:#fff;
                border:none; border-radius:4px; padding:6px 22px;
            }
            QPushButton:hover   { background:#333; }
            QPushButton:pressed { background:#000; }
            QPushButton:disabled { background:#DDD; color:#AAA; }
        """)
    else:
        btn.setStyleSheet("""
            QPushButton {
                background:#fff; color:#111;
                border:1px solid #CCC; border-radius:4px; padding:5px 14px;
            }
            QPushButton:hover   { border-color:#555; }
            QPushButton:pressed { background:#EEE; }
            QPushButton:disabled { color:#CCC; border-color:#EEE; }
        """)
    return btn


class CompareWindow(QMainWindow):

    def __init__(self):
        super().__init__()
        self.setWindowTitle("Comparer")
        self.setMinimumSize(1100, 680)
        self.setWindowIcon(_load_icon())
        self._syncing      = False
        self._result: Optional[dict] = None
        self._worker: Optional[CompareWorker] = None
        self._current_diff = 0
        self._diff_count   = 0
        self._sel_idx      = -1
        self._setup_ui()
        self._setup_shortcuts()

    def _setup_shortcuts(self):
        QShortcut(QKeySequence("Escape"), self, self._go_back)
        QShortcut(QKeySequence("Ctrl+E"), self, self._export_report)
        QShortcut(QKeySequence("Ctrl+P"), self, self._export_pdf)
        QShortcut(QKeySequence("Right"),  self, self._next_diff)
        QShortcut(QKeySequence("Left"),   self, self._prev_diff)

    def _setup_ui(self):
        self.stack = QStackedWidget()
        self.setCentralWidget(self.stack)
        self.stack.addWidget(self._build_welcome())
        self.stack.addWidget(self._build_compare())
        self.stack.setCurrentIndex(0)

    # ── Pagina 0: Welcome ─────────────────────────────────────────────────────

    def _build_welcome(self) -> QWidget:
        page = QWidget()
        page.setStyleSheet("background:#fff;")

        outer = QVBoxLayout(page)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.addStretch(2)

        center = QWidget()
        center.setMaximumWidth(820)
        center.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        center_layout = QVBoxLayout(center)
        center_layout.setContentsMargins(0, 0, 0, 0)
        center_layout.setSpacing(0)

        title = QLabel("Comparer")
        title.setAlignment(Qt.AlignCenter)
        f = QFont()
        f.setPointSize(28)
        f.setWeight(QFont.Light)
        title.setFont(f)
        title.setStyleSheet("color:#111;")

        subtitle = QLabel("Compara documente la nivel de paragraf (PDF cu PDF, DOCX cu DOCX)")
        subtitle.setAlignment(Qt.AlignCenter)
        fs = QFont()
        fs.setPointSize(9)
        subtitle.setFont(fs)
        subtitle.setStyleSheet("color:#AAA; margin-top:6px; margin-bottom:36px;")

        center_layout.addWidget(title)
        center_layout.addWidget(subtitle)

        zones = QHBoxLayout()
        zones.setSpacing(16)
        self.drop_left  = DropZone("Baza  —  Document 1")
        self.drop_right = DropZone("Comparatie  —  Document 2")
        self.drop_left.setFixedHeight(148)
        self.drop_right.setFixedHeight(148)
        self.drop_left.file_loaded.connect(self._on_file_loaded)
        self.drop_right.file_loaded.connect(self._on_file_loaded)
        zones.addWidget(self.drop_left)
        zones.addWidget(self.drop_right)
        center_layout.addLayout(zones)
        center_layout.addSpacing(16)

        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 0)
        self.progress_bar.setFixedHeight(3)
        self.progress_bar.setTextVisible(False)
        self.progress_bar.setStyleSheet("""
            QProgressBar { background:#F0F0F0; border:none; border-radius:1px; }
            QProgressBar::chunk { background:#111; border-radius:1px; }
        """)
        self.progress_bar.hide()
        center_layout.addWidget(self.progress_bar)
        center_layout.addSpacing(8)

        self.status_lbl = QLabel("")
        self.status_lbl.setAlignment(Qt.AlignCenter)
        fst = QFont()
        fst.setPointSize(9)
        self.status_lbl.setFont(fst)
        self.status_lbl.setStyleSheet("color:#888;")
        self.status_lbl.setFixedHeight(18)

        self.compare_btn = _flat_btn("Compara", dark=True)
        fb = QFont()
        fb.setPointSize(10)
        self.compare_btn.setFont(fb)
        self.compare_btn.setFixedHeight(40)
        self.compare_btn.setMinimumWidth(200)
        self.compare_btn.setEnabled(False)
        self.compare_btn.clicked.connect(lambda: self._run_comparison())

        btn_row = QHBoxLayout()
        btn_row.addStretch()
        btn_row.addWidget(self.compare_btn)
        btn_row.addStretch()

        center_layout.addWidget(self.status_lbl)
        center_layout.addSpacing(10)
        center_layout.addLayout(btn_row)

        h = QHBoxLayout()
        h.addStretch()
        h.addWidget(center)
        h.addStretch()
        outer.addLayout(h)
        outer.addStretch(3)

        return page

    # ── Pagina 1: Comparatie ──────────────────────────────────────────────────

    def _build_compare(self) -> QWidget:
        page = QWidget()
        page.setStyleSheet("background:#fff;")

        layout = QVBoxLayout(page)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        layout.addWidget(self._build_toolbar())

        self.splitter = QSplitter(Qt.Horizontal)
        self.splitter.setStyleSheet(
            "QSplitter::handle { background:#DDDDDD; width:1px; }"
        )

        left_panel_w = QWidget()
        ll = QVBoxLayout(left_panel_w)
        ll.setContentsMargins(0, 0, 0, 0)
        ll.setSpacing(0)
        self.left_hdr = QLabel("  Document 1")
        self.left_hdr.setFixedHeight(28)
        self.left_hdr.setStyleSheet(
            "background:#111;color:#fff;font-size:10px;letter-spacing:.4px;"
            "font-family:system-ui,Arial,sans-serif;"
        )
        self.left_hdr.setAlignment(Qt.AlignVCenter | Qt.AlignLeft)
        self.left_panel = ComparePanel()
        ll.addWidget(self.left_hdr)
        ll.addWidget(self.left_panel)

        right_panel_w = QWidget()
        rl = QVBoxLayout(right_panel_w)
        rl.setContentsMargins(0, 0, 0, 0)
        rl.setSpacing(0)
        self.right_hdr = QLabel("  Document 2")
        self.right_hdr.setFixedHeight(28)
        self.right_hdr.setStyleSheet(
            "background:#444;color:#fff;font-size:10px;letter-spacing:.4px;"
            "font-family:system-ui,Arial,sans-serif;"
        )
        self.right_hdr.setAlignment(Qt.AlignVCenter | Qt.AlignLeft)
        self.right_panel = ComparePanel()
        rl.addWidget(self.right_hdr)
        rl.addWidget(self.right_panel)

        self.splitter.addWidget(left_panel_w)
        self.splitter.addWidget(right_panel_w)
        self.splitter.setSizes([500, 500])
        layout.addWidget(self.splitter, 1)

        layout.addWidget(self._build_stats_panel())

        self.left_panel.block_clicked.connect(self._on_block_clicked)
        self.right_panel.block_clicked.connect(self._on_block_clicked)
        self.left_panel.verticalScrollBar().valueChanged.connect(self._sync_l2r)
        self.right_panel.verticalScrollBar().valueChanged.connect(self._sync_r2l)

        return page

    def _build_toolbar(self) -> QFrame:
        bar = QFrame()
        bar.setFixedHeight(46)
        bar.setStyleSheet("background:#fff; border-bottom:1px solid #E0E0E0;")

        lay = QHBoxLayout(bar)
        lay.setContentsMargins(12, 0, 12, 0)
        lay.setSpacing(8)

        back_btn        = _flat_btn("\u2190 Noua comparatie")
        self.swap_btn   = _flat_btn("\u21c4  Inverseaza")
        self.export_btn = _flat_btn("Exporta HTML")
        self.export_pdf_btn = _flat_btn("Exporta PDF")
        back_btn.clicked.connect(self._go_back)
        self.swap_btn.clicked.connect(self._swap_and_recompare)
        self.export_btn.clicked.connect(self._export_report)
        self.export_pdf_btn.clicked.connect(self._export_pdf)

        self.prev_diff_btn = _flat_btn("\u2039 Diff")
        self.next_diff_btn = _flat_btn("Diff \u203a")
        self.diff_counter_lbl = QLabel("\u2014")
        fc = QFont()
        fc.setPointSize(9)
        self.diff_counter_lbl.setFont(fc)
        self.diff_counter_lbl.setStyleSheet("color:#666; min-width:52px;")
        self.diff_counter_lbl.setAlignment(Qt.AlignCenter)
        self.prev_diff_btn.clicked.connect(self._prev_diff)
        self.next_diff_btn.clicked.connect(self._next_diff)

        self.left_name_lbl = QLabel("")
        f = QFont()
        f.setPointSize(9)
        self.left_name_lbl.setFont(f)
        self.left_name_lbl.setStyleSheet("color:#555;")

        self.right_name_lbl = QLabel("")
        self.right_name_lbl.setFont(f)
        self.right_name_lbl.setStyleSheet("color:#555;")

        sim_lbl = QLabel("Similaritate")
        sim_lbl.setFont(f)
        sim_lbl.setStyleSheet("color:#999;")

        self.sim_badge = QLabel("\u2014")
        fb = QFont()
        fb.setPointSize(11)
        fb.setWeight(QFont.DemiBold)
        self.sim_badge.setFont(fb)
        self.sim_badge.setStyleSheet(
            "color:#111; border:1px solid #CCC; border-radius:3px; padding:1px 10px;"
        )
        self.sim_badge.setAlignment(Qt.AlignCenter)
        self.sim_badge.setMinimumWidth(72)

        lay.addWidget(back_btn)
        lay.addWidget(self.swap_btn)
        lay.addSpacing(8)
        lay.addWidget(self.prev_diff_btn)
        lay.addWidget(self.diff_counter_lbl)
        lay.addWidget(self.next_diff_btn)
        lay.addSpacing(8)
        lay.addWidget(self.left_name_lbl)
        lay.addStretch()
        lay.addWidget(sim_lbl)
        lay.addSpacing(4)
        lay.addWidget(self.sim_badge)
        lay.addStretch()
        lay.addWidget(self.right_name_lbl)
        lay.addSpacing(12)
        lay.addWidget(self.export_btn)
        lay.addWidget(self.export_pdf_btn)

        return bar

    def _build_stats_panel(self) -> QFrame:
        panel = QFrame()
        panel.setFixedHeight(76)
        panel.setStyleSheet("background:#F7F7F7; border-top:1px solid #E0E0E0;")

        lay = QHBoxLayout(panel)
        lay.setContentsMargins(20, 0, 20, 0)
        lay.setSpacing(0)

        sim_w = QWidget()
        sim_w.setStyleSheet("background:transparent;")
        sl = QVBoxLayout(sim_w)
        sl.setSpacing(0)
        sl.setContentsMargins(0, 0, 20, 0)
        self.sim_big = QLabel("\u2014")
        fb = QFont()
        fb.setPointSize(26)
        fb.setWeight(QFont.Light)
        self.sim_big.setFont(fb)
        self.sim_big.setStyleSheet("color:#111;")
        self.sim_big.setAlignment(Qt.AlignCenter)
        sl_lbl = QLabel("SIMILARITATE")
        fs = QFont()
        fs.setPointSize(7)
        sl_lbl.setFont(fs)
        sl_lbl.setStyleSheet("color:#AAA; letter-spacing:1px;")
        sl_lbl.setAlignment(Qt.AlignCenter)
        sl.addWidget(self.sim_big)
        sl.addWidget(sl_lbl)

        sep = QFrame()
        sep.setFrameShape(QFrame.VLine)
        sep.setStyleSheet("color:#DDDDDD;")
        sep.setFixedWidth(1)

        (self.s_equal,   w_eq)  = self._make_stat("IDENTICE",   "\u2014")
        (self.s_mod,     w_mod) = self._make_stat("MODIFICATE", "\u2014")
        (self.s_added,   w_add) = self._make_stat("ADAUGATE",   "\u2014")
        (self.s_deleted, w_del) = self._make_stat("STERSE",     "\u2014")
        (self.s_tot1,    w_t1)  = self._make_stat("PAR. DOC 1", "\u2014")
        (self.s_tot2,    w_t2)  = self._make_stat("PAR. DOC 2", "\u2014")

        legend = self._build_legend()

        lay.addWidget(sim_w)
        lay.addWidget(sep)
        lay.addSpacing(4)
        for w in [w_eq, w_mod, w_add, w_del, w_t1, w_t2]:
            lay.addWidget(w)
        lay.addStretch()
        lay.addWidget(legend)

        return panel

    def _make_stat(self, label: str, value: str):
        w = QWidget()
        w.setStyleSheet("background:transparent;")
        lay = QVBoxLayout(w)
        lay.setSpacing(0)
        lay.setContentsMargins(14, 8, 14, 8)

        val = QLabel(value)
        fv = QFont()
        fv.setPointSize(15)
        fv.setWeight(QFont.DemiBold)
        val.setFont(fv)
        val.setStyleSheet("color:#111;")
        val.setAlignment(Qt.AlignCenter)

        lbl = QLabel(label)
        fl = QFont()
        fl.setPointSize(7)
        lbl.setFont(fl)
        lbl.setStyleSheet("color:#AAA; letter-spacing:.5px;")
        lbl.setAlignment(Qt.AlignCenter)

        lay.addWidget(val)
        lay.addWidget(lbl)
        return val, w

    def _build_legend(self) -> QWidget:
        w = QWidget()
        w.setStyleSheet("background:transparent;")
        lay = QVBoxLayout(w)
        lay.setSpacing(5)
        lay.setContentsMargins(8, 8, 0, 8)

        items = [
            ("#FFCDD2", "#E53935", "Sters  (doc. stanga)"),
            ("#C8E6C9", "#43A047", "Adaugat  (doc. dreapta)"),
            ("#FFECB3", "#FFB300", "Modificat  (ambele)"),
        ]
        fl = QFont()
        fl.setPointSize(8)
        for border, bg, txt in items:
            row = QHBoxLayout()
            row.setSpacing(7)
            swatch = QLabel()
            swatch.setFixedSize(22, 13)
            swatch.setStyleSheet(
                f"background:{bg};border-left:3px solid {border};border-radius:1px;"
            )
            lbl = QLabel(txt)
            lbl.setFont(fl)
            lbl.setStyleSheet("color:#888;")
            row.addWidget(swatch)
            row.addWidget(lbl)
            lay.addLayout(row)
        return w

    # ── Logica aplicatiei ────────────────────────────────────────────────────

    def _on_file_loaded(self, _path: str):
        both = self.drop_left.file_path and self.drop_right.file_path
        self.compare_btn.setEnabled(bool(both))
        # Show inline warning if types mismatch
        if both:
            err = validate_same_type(self.drop_left.file_path, self.drop_right.file_path)
            if err:
                self.status_lbl.setText(
                    "\u26a0 Tipuri diferite! Comparatia e permisa doar PDF\u2194PDF sau DOCX\u2194DOCX."
                )
                self.status_lbl.setStyleSheet("color:#E53935;")
                self.compare_btn.setEnabled(False)
            else:
                self.status_lbl.setText("")
                self.status_lbl.setStyleSheet("color:#888;")
        else:
            self.status_lbl.setText("")
            self.status_lbl.setStyleSheet("color:#888;")

    def _set_busy(self, busy: bool):
        self.compare_btn.setEnabled(not busy and bool(
            self.drop_left.file_path and self.drop_right.file_path
        ))
        self.swap_btn.setEnabled(not busy)
        self.export_btn.setEnabled(not busy and self._result is not None)
        self.export_pdf_btn.setEnabled(not busy and self._result is not None)
        if busy:
            self.progress_bar.show()
        else:
            self.progress_bar.hide()
            self.status_lbl.setText("")

    def _run_comparison(self, path1: str = None, path2: str = None):
        p1 = path1 or self.drop_left.file_path
        p2 = path2 or self.drop_right.file_path
        if not p1 or not p2:
            return

        # Validate same file type
        err = validate_same_type(p1, p2)
        if err:
            QMessageBox.warning(self, "Tip incompatibil", err)
            return

        if self._worker and self._worker.isRunning():
            try:
                self._worker.finished.disconnect()
                self._worker.error.disconnect()
            except TypeError:
                pass

        self._set_busy(True)
        self.status_lbl.setText("Se proceseaza...")

        self._worker = CompareWorker(p1, p2)
        self._worker.finished.connect(self._on_comparison_done)
        self._worker.error.connect(self._on_comparison_error)
        self._worker.progress.connect(self.status_lbl.setText)
        self._worker.start()

    @pyqtSlot(dict)
    def _on_comparison_done(self, result: dict):
        self._result = result
        self._set_busy(False)
        self._update_compare_page(result)
        self.stack.setCurrentIndex(1)

    @pyqtSlot(str)
    def _on_comparison_error(self, msg: str):
        self._set_busy(False)
        QMessageBox.critical(self, "Eroare", msg)

    def _on_block_clicked(self, idx: int):
        new_sel = -1 if self._sel_idx == idx else idx
        self._sel_idx = new_sel
        self.left_panel.set_selected(new_sel)
        self.right_panel.set_selected(new_sel)

    def _update_compare_page(self, result: dict):
        self._sel_idx = -1
        sim    = result["similarity"]
        stats  = result["stats"]
        sim_pct = f"{sim * 100:.1f}%"

        self.sim_big.setText(sim_pct)
        self.sim_badge.setText(sim_pct)
        self.s_equal.setText(str(stats["equal"]))
        self.s_mod.setText(str(stats["modified"]))
        self.s_added.setText(str(stats["added"]))
        self.s_deleted.setText(str(stats["deleted"]))
        self.s_tot1.setText(str(stats["total_left"]))
        self.s_tot2.setText(str(stats["total_right"]))

        n1 = Path(result["path1"]).name
        n2 = Path(result["path2"]).name
        self.left_name_lbl.setText(n1)
        self.right_name_lbl.setText(n2)
        self.left_hdr.setText(f"  {n1}")
        self.right_hdr.setText(f"  {n2}")

        self._diff_count   = result.get("diff_count", 0)
        self._current_diff = 0
        self._update_diff_nav()

        self.left_panel.load_blocks(
            result["left_blocks"], "left", result["diff_positions"]
        )
        self.right_panel.load_blocks(
            result["right_blocks"], "right", result["diff_positions"]
        )

    def _update_diff_nav(self):
        has = self._diff_count > 0
        self.prev_diff_btn.setEnabled(has and self._current_diff > 0)
        self.next_diff_btn.setEnabled(has and self._current_diff < self._diff_count - 1)
        self.diff_counter_lbl.setText(
            f"{self._current_diff + 1} / {self._diff_count}" if has else "0 / 0"
        )

    def _go_to_diff(self, idx: int):
        if not self._diff_count:
            return
        self._current_diff = max(0, min(idx, self._diff_count - 1))
        block_idx = self._result["diff_positions"][self._current_diff]
        self._syncing = True
        self.left_panel.scroll_to_block(block_idx)
        self.right_panel.scroll_to_block(block_idx)
        self._syncing = False
        self._update_diff_nav()

    def _next_diff(self):
        if self.stack.currentIndex() == 1:
            self._go_to_diff(self._current_diff + 1)

    def _prev_diff(self):
        if self.stack.currentIndex() == 1:
            self._go_to_diff(self._current_diff - 1)

    def _go_back(self):
        if self.stack.currentIndex() == 0:
            return
        if self._worker and self._worker.isRunning():
            try:
                self._worker.finished.disconnect()
                self._worker.error.disconnect()
            except TypeError:
                pass
        self._set_busy(False)
        self._sel_idx = -1
        self.stack.setCurrentIndex(0)
        self.drop_left.reset()
        self.drop_right.reset()
        self.compare_btn.setEnabled(False)
        self._result = None

    def _swap_and_recompare(self):
        if not self._result:
            return
        p1 = self._result["path1"]
        p2 = self._result["path2"]
        self.drop_left._set_file_display(p2)
        self.drop_right._set_file_display(p1)
        self._run_comparison(p2, p1)

    def _export_report(self):
        if not self._result:
            return
        now = datetime.now()
        default = now.strftime("%Y%m%d") + "_N_xxx_RaportComparatie-[xxxxx].html"
        path, _ = QFileDialog.getSaveFileName(
            self, "Salveaza raport HTML", default, "HTML (*.html)"
        )
        if not path:
            return
        try:
            content = export_html_report(
                self._result["path1"], self._result["path2"], self._result,
            )
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
            QMessageBox.information(self, "Export reusit", f"Salvat:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Eroare la export", str(e))

    def _export_pdf(self):
        """Export comparison report as PDF using QPrinter."""
        if not self._result:
            return
        now = datetime.now()
        default = now.strftime("%Y%m%d") + "_N_xxx_RaportComparatie-[xxxxx].pdf"
        path, _ = QFileDialog.getSaveFileName(
            self, "Salveaza raport PDF", default, "PDF (*.pdf)"
        )
        if not path:
            return
        try:
            html_content = export_pdf_report(
                self._result["path1"], self._result["path2"], self._result,
            )

            printer = QPrinter(QPrinter.ScreenResolution)
            printer.setOutputFormat(QPrinter.PdfFormat)
            printer.setOutputFileName(path)
            printer.setPageSize(QPrinter.A4)
            printer.setPageMargins(12, 12, 12, 12, QPrinter.Millimeter)

            doc = QTextDocument()
            doc.setHtml(html_content)
            # Set page size based on printer's printable area in device pixels
            page_rect = printer.pageRect(QPrinter.Point)
            doc.setPageSize(QRectF(0, 0, page_rect.width(), page_rect.height()).size())
            doc.print_(printer)

            QMessageBox.information(self, "Export PDF reusit", f"Salvat:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Eroare la export PDF", str(e))

    def _sync_l2r(self, value: int):
        if not self._syncing:
            self._syncing = True
            self.right_panel.verticalScrollBar().setValue(value)
            self._syncing = False

    def _sync_r2l(self, value: int):
        if not self._syncing:
            self._syncing = True
            self.left_panel.verticalScrollBar().setValue(value)
            self._syncing = False


# ══════════════════════════════════════════════════════════════════════════════
# 7. Entry point
# ══════════════════════════════════════════════════════════════════════════════

def main():
    QApplication.setAttribute(Qt.AA_EnableHighDpiScaling, True)
    QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps, True)

    app = QApplication(sys.argv)
    app.setApplicationName("Comparer")
    app.setStyle("Fusion")
    app.setWindowIcon(_load_icon())

    palette = QPalette()
    palette.setColor(QPalette.Window,          QColor("#FFFFFF"))
    palette.setColor(QPalette.WindowText,      QColor("#111111"))
    palette.setColor(QPalette.Base,            QColor("#FFFFFF"))
    palette.setColor(QPalette.AlternateBase,   QColor("#F7F7F7"))
    palette.setColor(QPalette.Text,            QColor("#111111"))
    palette.setColor(QPalette.ButtonText,      QColor("#111111"))
    palette.setColor(QPalette.Button,          QColor("#FFFFFF"))
    palette.setColor(QPalette.Highlight,       QColor("#111111"))
    palette.setColor(QPalette.HighlightedText, QColor("#FFFFFF"))
    app.setPalette(palette)

    win = CompareWindow()
    win.showMaximized()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
